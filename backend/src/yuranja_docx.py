"""Export approved Yuranja exhibitions to a Word (.docx) review document."""

from __future__ import annotations

import sys
from pathlib import Path
from typing import Any

SRC_DIR = Path(__file__).resolve().parent
if str(SRC_DIR) not in sys.path:
    sys.path.insert(0, str(SRC_DIR))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

import database
import yuranja_candidates as yc
import yuranja_export as ye
import yuranja_model as ym

DOCX_PATH = database.ROOT_DATA_DIR / "yuranja_exhibitions_approved.docx"


def _fmt_dates(dates: dict[str, Any]) -> str:
    start = (dates or {}).get("start") or "—"
    end = (dates or {}).get("end") or "—"
    return f"{start} → {end}"


def _artists(record: dict[str, Any]) -> str:
    arts = record.get("artists") or []
    return ", ".join(arts) if arts else "—"


def _admission(record: dict[str, Any]) -> str:
    admission = record.get("admission") or {}
    status = admission.get("status") or "unknown"
    if status == "unknown":
        return "Admission not published"
    return str(admission.get("display") or "Admission not published")


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    """Add a labelled hyperlink run to a paragraph."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(u)
    new_run.append(r_pr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _citations(record: dict[str, Any]) -> tuple[str, str]:
    ex_url = ""
    ad_url = ""
    for c in record.get("citations") or []:
        if not isinstance(c, dict):
            continue
        if c.get("type") == "exhibition" and c.get("url"):
            ex_url = str(c["url"])
        if c.get("type") == "admission" and c.get("url"):
            ad_url = str(c["url"])
    if not ex_url:
        ex_url = str(record.get("website") or record.get("exhibitionUrl") or "")
    return ex_url, ad_url


def _load_approved_records() -> list[dict[str, Any]]:
    conn = database.connect()
    database.init_schema(conn)
    rows = conn.execute(
        """
        SELECT *
        FROM exhibitions
        WHERE lower(COALESCE(editorial_status, 'pending')) = 'approved'
          AND lower(COALESCE(archive_status, 'active')) = 'active'
          AND COALESCE(is_duplicate, 0) = 0
          AND lower(COALESCE(status, '')) IN ('current', 'upcoming')
          AND COALESCE(trim(title), '') NOT IN ('', '(unavailable)')
          AND COALESCE(trim(start_date), '') != ''
          AND COALESCE(trim(end_date), '') != ''
        ORDER BY city, name, start_date, title
        """
    ).fetchall()

    records: list[dict[str, Any]] = []
    for row in rows:
        record = ye._row_to_record(row)
        # Prefer cleaned title / artists from candidate enrichment when available.
        enriched = yc.enrich_candidate_fields(
            {
                **record,
                "exhibitionUrl": record.get("exhibitionUrl") or record.get("website") or "",
            }
        )
        record["title"] = enriched.get("title") or record.get("title")
        if enriched.get("artists"):
            record["artists"] = enriched["artists"]
        if enriched.get("format"):
            record["format"] = enriched["format"]
        if not record.get("description"):
            record["description"] = yc.build_description(record)
        if ym.export_eligible(record) or (
            record.get("dates", {}).get("start") and record.get("dates", {}).get("end")
        ):
            records.append(record)
    return records


def export_yuranja_docx(*, path: Path | None = None) -> dict[str, Any]:
    out = path or DOCX_PATH
    records = _load_approved_records()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Yuranja — Approved exhibitions", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    intro = doc.add_paragraph()
    intro.add_run(
        f"Editorial export of {len(records)} approved, current or upcoming exhibitions "
        f"with verified dates. Generated {database.now_iso()[:10]}."
    )

    # Summary by city
    doc.add_heading("Summary by city", level=1)
    by_city: dict[str, int] = {}
    for r in records:
        by_city[r.get("city") or "Unknown"] = by_city.get(r.get("city") or "Unknown", 0) + 1
    for city in sorted(by_city):
        doc.add_paragraph(f"{city}: {by_city[city]}", style="List Bullet")

    current_city = None
    for record in records:
        city = record.get("city") or "Unknown"
        if city != current_city:
            doc.add_heading(city, level=1)
            current_city = city

        doc.add_heading(str(record.get("title") or "Untitled"), level=2)

        meta = doc.add_paragraph()
        meta.add_run("Institution: ").bold = True
        meta.add_run(f"{record.get('venue') or '—'}")
        if record.get("country"):
            meta.add_run(f" · {record.get('country')}")

        p = doc.add_paragraph()
        p.add_run("Artists: ").bold = True
        p.add_run(_artists(record))

        p = doc.add_paragraph()
        p.add_run("Dates: ").bold = True
        p.add_run(_fmt_dates(record.get("dates") or {}))
        p.add_run(f"  ·  {record.get('status') or ''}")

        if record.get("format"):
            p = doc.add_paragraph()
            p.add_run("Format: ").bold = True
            p.add_run(str(record.get("format")))

        if record.get("address"):
            p = doc.add_paragraph()
            p.add_run("Address: ").bold = True
            p.add_run(str(record.get("address")))

        p = doc.add_paragraph()
        p.add_run("Admission: ").bold = True
        admission = record.get("admission") or {}
        if (admission.get("status") or "unknown") == "unknown":
            p.add_run("Admission not published")
            link_p = doc.add_paragraph()
            label = str(admission.get("informationLabel") or "Official visitor information")
            url = str(admission.get("informationUrl") or admission.get("ticketUrl") or "")
            if url:
                _add_hyperlink(link_p, url, label)
            else:
                link_p.add_run("No official admission information URL available").italic = True
        else:
            p.add_run(_admission(record))
            ticket = str(admission.get("ticketUrl") or "")
            if ticket:
                link_p = doc.add_paragraph()
                _add_hyperlink(link_p, ticket, "Official admission page")

        if record.get("description"):
            p = doc.add_paragraph()
            p.add_run("Description: ").bold = True
            p.add_run(str(record.get("description")))

        note = str(record.get("yuranjaNote") or "").strip()
        p = doc.add_paragraph()
        p.add_run("Yuranja note: ").bold = True
        run = p.add_run(note if note else "(empty — human editorial note not yet written)")
        if not note:
            run.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        ex_url, ad_url = _citations(record)
        p = doc.add_paragraph()
        p.add_run("Official exhibition URL: ").bold = True
        if ex_url:
            _add_hyperlink(p, ex_url, "Official exhibition page")
        else:
            p.add_run("—")

        p = doc.add_paragraph()
        p.add_run("Admission check: ").bold = True
        info_url = str(admission.get("informationUrl") or admission.get("ticketUrl") or ad_url or "")
        info_label = str(admission.get("informationLabel") or "Official visitor information")
        if info_url:
            _add_hyperlink(p, info_url, info_label)
        else:
            p.add_run("none verified")

        p = doc.add_paragraph()
        p.add_run("Slug: ").bold = True
        p.add_run(str(record.get("slug") or "—"))

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    msg = f"Exported {len(records)} approved exhibitions to {out}"
    print(msg)
    return {"success": True, "message": msg, "path": str(out), "count": len(records)}


if __name__ == "__main__":
    export_yuranja_docx()
